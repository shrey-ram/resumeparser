import jinja2
import re
from docx import Document
from dateparser import parse

def is_contact_info(text):
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    phone_pattern = r'\+?\d[\d -]{7,}\d'
    return re.search(email_pattern, text) or re.search(phone_pattern, text)

def extract_dates(text):
    dates = []
    for match in re.finditer(r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s\d{4}\b', text):
        date_str = match.group()
        if date_str not in dates:
            date = parse(date_str)
            if date:
                dates.append(date.strftime("%b %Y"))  # Format date as "Mon YYYY"
    return dates

def parse_docx(file_path):
    doc = Document(file_path)

    name = doc.paragraphs[0].text.strip()
    contact_info = []
    sections = {
        "Professional Experience": [],
        "Skills": [],
        "Education": [],
        "Certifications": [],
        "Projects": [],
    }
    current_section = None
    current_workplace = None
    current_points = []

    section_keywords = {
        "professional experience": "Professional Experience",
        "skills": "Skills",
        "education": "Education",
        "certifications": "Certifications",
        "projects": "Projects"
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if not text:
            continue

        # Detect contact information
        if is_contact_info(text):
            contact_info.append(text)
            continue

        # Detect sections by matching keywords
        lower_text = text.lower()
        for key in section_keywords:
            if key in lower_text:
                if current_section and current_workplace:
                    sections[current_section].append({"workplace": current_workplace, "points": current_points})
                    current_points = []
                current_section = section_keywords[key]
                current_workplace = None
                break
        else:
            if current_section is None:
                continue  # Skip paragraph if it's not part of any section

            # Adjust handling for different sections
            if current_section in ["Professional Experience", "Education", "Projects"]:
                # Detect workplace or project name
                if current_workplace is None:
                    current_workplace = text
                else:
                    current_points.append(text)
            elif current_section == "Certifications":
                sections[current_section].append(text)
            elif current_section == "Skills":
                sections[current_section].append(text)

    # Final append of last section data
    if current_section and current_workplace:
        sections[current_section].append({"workplace": current_workplace, "points": current_points})

    return {
        "name": name,
        "contact_info": contact_info,
        "sections": sections
    }

def create_html(parsed_content):
    template_loader = jinja2.FileSystemLoader('.')
    template_env = jinja2.Environment(loader=template_loader)

    template = template_env.get_template('base.html')
    html = template.render(
        name=parsed_content["name"],
        contact_info=parsed_content["contact_info"],
        sections=parsed_content["sections"]
    )

    with open('/mnt/f/resume1.html', 'w') as f:
        f.write(html)

def print_section(title, content):
    print(f"\n{title}:")
    if isinstance(content, list):
        for item in content:
            if isinstance(item, dict):
                print(f"\nWorkplace: {item['workplace']}")
                for point in item['points']:
                    print(f"- {point}")
            else:
                print(f"- {item}")
    else:
        print(content)

def main():
    file_path = r"/mnt/f/Shreyas Cybersecurity resume.docx"
    parsed_content = parse_docx(file_path)

    print("Name:")
    print(parsed_content["name"])

    print("\nContact Information:")
    for contact in parsed_content["contact_info"]:
        print(contact)

    print_section("Professional Experience", parsed_content["sections"]["Professional Experience"])
    print_section("Skills", parsed_content["sections"]["Skills"])
    print_section("Education", parsed_content["sections"]["Education"])
    print_section("Certifications", parsed_content["sections"]["Certifications"])
    print_section("Projects", parsed_content["sections"]["Projects"])

    create_html(parsed_content)

if __name__ == "__main__":
    main()
