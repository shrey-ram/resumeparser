import jinja2
import re
from docx import Document
from dateparser import parse
import os

def is_contact_info(text):
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    phone_pattern = r'\+?\d[\d -]{7,}\d'
    return re.search(email_pattern, text) or re.search(phone_pattern, text)

def extract_dates(text):
    dates = []
    for match in re.finditer(r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s\d{4}\b', text):
        date_str = match.group()
        if date_str not in dates:
            date = parse(date_str)
            if date:
                dates.append(date.strftime("%b %Y"))  # Format date as "Mon YYYY"
    return dates

def parse_docx(file_path):
    doc = Document(file_path)

    name = ""
    contact_info = []
    sections = {
        "Professional Experience": [],
        "Skills": [],
        "Education": [],
        "Certifications": [],
        "Projects": [],
    }
    current_section = None
    current_workplace = None
    current_points = []

    section_keywords = {
        "professional experience": "Professional Experience",
        "skills": "Skills",
        "education": "Education",
        "certifications": "Certifications",
        "projects": "Projects"
    }

    def is_workplace(text):
        # Define a regex pattern to detect workplace entries more reliably
        return re.match(r'.*,\s.*\s\d{4} - \w*', text) or "Engineer" in text or "Developer" in text or "Assistant" in text or "University" in text or "Project" in text

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if not text:
            continue

        if is_contact_info(text):
            contact_info.append(text)
        elif text.lower() in section_keywords:
            if current_section and current_workplace:
                if current_section in ["Professional Experience", "Education", "Certifications", "Projects"]:
                    sections[current_section].append({"workplace": current_workplace, "points": current_points})
                current_points = []
            current_section = section_keywords[text.lower()]
            current_workplace = None
        else:
            if current_section is None:
                continue  # Skip paragraph if it's not part of any section

            # Handle new workplace detection
            if current_section in ["Professional Experience", "Education", "Certifications", "Projects"]:
                if is_workplace(text):
                    if current_workplace:
                        # Save the previous workplace's points
                        sections[current_section].append({"workplace": current_workplace, "points": current_points})
                    # Start a new workplace
                    current_workplace = text
                    current_points = []
                else:
                    if current_section == "Certifications":
                        sections[current_section].append(text)
                    else:
                        current_points.append(text)
            else:
                sections[current_section].append(text)

    if current_section and current_workplace:
        if current_section in ["Professional Experience", "Education", "Certifications", "Projects"]:
            sections[current_section].append({"workplace": current_workplace, "points": current_points})
    elif current_section:
        if current_section in ["Professional Experience", "Education", "Certifications", "Projects"]:
            sections[current_section].append({"workplace": "", "points": current_points})

    return {
        "name": doc.paragraphs[0].text.strip(),  # Extract name from first paragraph
        "contact_info": contact_info,
        "sections": sections
    }


def create_html(parsed_content, output_path):
    template_loader = jinja2.FileSystemLoader('.')
    template_env = jinja2.Environment(loader=template_loader)

    template = template_env.get_template('base.html')
    html = template.render(
        name=parsed_content["name"],
        contact_info=parsed_content["contact_info"],
        sections=parsed_content["sections"]
    )

    with open(output_path, 'w') as f:
        f.write(html)


    # Save the HTML to a file
    html_file_path = 'resume1.html'
    with open(html_file_path, 'w') as f:
        f.write(html)

    return html_file_path
